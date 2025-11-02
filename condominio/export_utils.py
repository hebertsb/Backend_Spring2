"""
Utilidades para exportación de reportes a PDF y Excel.
"""
from reportlab.lib.pagesizes import letter, A4
from reportlab.lib import colors
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer, PageBreak
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.enums import TA_CENTER, TA_LEFT
from io import BytesIO
from datetime import datetime
from decimal import Decimal

from openpyxl import Workbook
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
from openpyxl.utils import get_column_letter
from openpyxl.worksheet.worksheet import Worksheet
from typing import cast

# Word/DOCX generation
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml


class ExportadorReportesPDF:
    """
    Clase para generar reportes en formato PDF con estructura profesional.
    """
    
    def __init__(self, moneda='USD'):
        self.styles = getSampleStyleSheet()
        self.moneda = moneda  # USD o BOB
        self.simbolo_moneda = '$' if moneda == 'USD' else 'Bs.'
        self.tasa_cambio = 6.96  # 1 USD = 6.96 BOB
        self._crear_estilos_personalizados()
    
    def _crear_estilos_personalizados(self):
        """Crea estilos personalizados para el documento."""
        # Estilo para título principal
        self.styles.add(ParagraphStyle(
            name='TituloReporte',
            parent=self.styles['Title'],
            fontSize=20,
            textColor=colors.HexColor('#2C3E50'),
            spaceAfter=30,
            alignment=TA_CENTER,
            fontName='Helvetica-Bold'
        ))
        
        # Estilo para subtítulos
        self.styles.add(ParagraphStyle(
            name='SubtituloReporte',
            parent=self.styles['Heading2'],
            fontSize=14,
            textColor=colors.HexColor('#34495E'),
            spaceAfter=12,
            spaceBefore=12,
            fontName='Helvetica-Bold'
        ))
        
        # Estilo para información
        self.styles.add(ParagraphStyle(
            name='InfoReporte',
            parent=self.styles['Normal'],
            fontSize=10,
            textColor=colors.HexColor('#7F8C8D'),
            spaceAfter=6
        ))
    
    def _formatear_valor(self, valor, convertir_moneda=True):
        """
        Formatea valores para mostrar en PDF con la moneda correcta.
        
        Args:
            valor: El valor a formatear
            convertir_moneda: Si True, convierte valores monetarios según self.moneda
        """
        if valor is None:
            return 'N/A'
        
        if isinstance(valor, Decimal):
            valor_float = float(valor)
            # Si el valor viene en BOB y queremos USD, convertir
            if convertir_moneda and self.moneda == 'USD':
                valor_float = valor_float / self.tasa_cambio
            # Si el valor viene en USD y queremos BOB, convertir
            elif convertir_moneda and self.moneda == 'BOB':
                # Asumir que los valores ya vienen en la moneda correcta
                pass
            return f"{self.simbolo_moneda} {valor_float:,.2f}"
        
        if isinstance(valor, (int, float)):
            valor_float = float(valor)
            if convertir_moneda and isinstance(valor, float):
                if self.moneda == 'USD':
                    valor_float = valor_float / self.tasa_cambio
            return f"{self.simbolo_moneda} {valor_float:,.2f}" if isinstance(valor, float) else f"{valor:,}"
        
        if isinstance(valor, datetime):
            return valor.strftime('%d/%m/%Y %H:%M')
        
        return str(valor)
    
    def generar_reporte_ventas_general(self, reporte_data, titulo="Reporte de Ventas General"):
        """
        Genera PDF para reporte de ventas general.
        """
        buffer = BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=letter, topMargin=0.75*inch, bottomMargin=0.75*inch)
        story = []
        
        # Título
        story.append(Paragraph(titulo, self.styles['TituloReporte']))
        story.append(Spacer(1, 0.2*inch))
        
        # Información del período
        periodo = reporte_data.get('periodo', {})
        if periodo.get('fecha_inicio') or periodo.get('fecha_fin'):
            fecha_inicio = periodo.get('fecha_inicio', 'N/A')
            fecha_fin = periodo.get('fecha_fin', 'N/A')
            if isinstance(fecha_inicio, datetime):
                fecha_inicio = fecha_inicio.strftime('%d/%m/%Y')
            if isinstance(fecha_fin, datetime):
                fecha_fin = fecha_fin.strftime('%d/%m/%Y')
            
            periodo_texto = f"<b>Período:</b> {fecha_inicio} - {fecha_fin}"
            story.append(Paragraph(periodo_texto, self.styles['InfoReporte']))
            story.append(Spacer(1, 0.1*inch))
        
        fecha_generacion = f"<b>Fecha de generación:</b> {datetime.now().strftime('%d/%m/%Y %H:%M')}"
        story.append(Paragraph(fecha_generacion, self.styles['InfoReporte']))
        story.append(Spacer(1, 0.3*inch))
        
        # Métricas Generales
        metricas = reporte_data.get('metricas_generales', {})
        if metricas:
            story.append(Paragraph("Métricas Generales", self.styles['SubtituloReporte']))
            
            data_metricas = [
                ['Métrica', 'Valor'],
                ['Total Ventas', self._formatear_valor(metricas.get('total_ventas'))],
                ['Cantidad de Reservas', self._formatear_valor(metricas.get('cantidad_reservas'))],
                ['Ticket Promedio', self._formatear_valor(metricas.get('ticket_promedio'))],
                ['Total Pagado', self._formatear_valor(metricas.get('total_pagado'))],
            ]
            
            table_metricas = Table(data_metricas, colWidths=[3*inch, 2.5*inch])
            table_metricas.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#3498DB')),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                ('ALIGN', (1, 1), (1, -1), 'RIGHT'),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, 0), 11),
                ('FONTSIZE', (0, 1), (-1, -1), 10),
                ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                ('TOPPADDING', (0, 1), (-1, -1), 8),
                ('BOTTOMPADDING', (0, 1), (-1, -1), 8),
                ('GRID', (0, 0), (-1, -1), 1, colors.HexColor('#BDC3C7')),
                ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, colors.HexColor('#ECF0F1')])
            ]))
            story.append(table_metricas)
            story.append(Spacer(1, 0.3*inch))
        
        # Top Paquetes
        top_paquetes = reporte_data.get('top_paquetes', [])
        if top_paquetes:
            story.append(Paragraph("Top Paquetes Más Vendidos", self.styles['SubtituloReporte']))
            
            data_paquetes = [['Paquete', 'Total Ventas', 'Cantidad']]
            for paquete in top_paquetes[:10]:
                data_paquetes.append([
                    paquete.get('paquete__nombre', 'N/A'),
                    self._formatear_valor(paquete.get('total_ventas')),
                    str(paquete.get('cantidad', 0))
                ])
            
            table_paquetes = Table(data_paquetes, colWidths=[3*inch, 1.5*inch, 1*inch])
            table_paquetes.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#27AE60')),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                ('ALIGN', (1, 1), (-1, -1), 'RIGHT'),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, 0), 11),
                ('FONTSIZE', (0, 1), (-1, -1), 9),
                ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                ('TOPPADDING', (0, 1), (-1, -1), 6),
                ('BOTTOMPADDING', (0, 1), (-1, -1), 6),
                ('GRID', (0, 0), (-1, -1), 1, colors.HexColor('#BDC3C7')),
                ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, colors.HexColor('#ECF0F1')])
            ]))
            story.append(table_paquetes)
            story.append(Spacer(1, 0.3*inch))
        
        # Top Servicios
        top_servicios = reporte_data.get('top_servicios', [])
        if top_servicios:
            story.append(Paragraph("Top Servicios Más Vendidos", self.styles['SubtituloReporte']))
            
            data_servicios = [['Servicio', 'Total Ventas', 'Cantidad']]
            for servicio in top_servicios[:10]:
                data_servicios.append([
                    servicio.get('servicio__titulo', 'N/A'),
                    self._formatear_valor(servicio.get('total_ventas')),
                    str(servicio.get('cantidad', 0))
                ])
            
            table_servicios = Table(data_servicios, colWidths=[3*inch, 1.5*inch, 1*inch])
            table_servicios.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#E67E22')),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                ('ALIGN', (1, 1), (-1, -1), 'RIGHT'),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, 0), 11),
                ('FONTSIZE', (0, 1), (-1, -1), 9),
                ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                ('TOPPADDING', (0, 1), (-1, -1), 6),
                ('BOTTOMPADDING', (0, 1), (-1, -1), 6),
                ('GRID', (0, 0), (-1, -1), 1, colors.HexColor('#BDC3C7')),
                ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, colors.HexColor('#ECF0F1')])
            ]))
            story.append(table_servicios)
            story.append(Spacer(1, 0.3*inch))
        
        # Top Clientes
        top_clientes = reporte_data.get('top_clientes', [])
        if top_clientes:
            story.append(PageBreak())
            story.append(Paragraph("Top Clientes", self.styles['SubtituloReporte']))
            
            data_clientes = [['Cliente', 'Total Gastado', 'Reservas']]
            for cliente in top_clientes[:15]:
                data_clientes.append([
                    cliente.get('cliente__nombre', 'N/A'),
                    self._formatear_valor(cliente.get('total_gastado')),
                    str(cliente.get('cantidad_reservas', 0))
                ])
            
            table_clientes = Table(data_clientes, colWidths=[3*inch, 1.5*inch, 1*inch])
            table_clientes.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#9B59B6')),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                ('ALIGN', (1, 1), (-1, -1), 'RIGHT'),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, 0), 11),
                ('FONTSIZE', (0, 1), (-1, -1), 9),
                ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                ('TOPPADDING', (0, 1), (-1, -1), 6),
                ('BOTTOMPADDING', (0, 1), (-1, -1), 6),
                ('GRID', (0, 0), (-1, -1), 1, colors.HexColor('#BDC3C7')),
                ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, colors.HexColor('#ECF0F1')])
            ]))
            story.append(table_clientes)
        
        # Construir PDF
        doc.build(story)
        buffer.seek(0)
        return buffer
    
    def generar_reporte_clientes(self, reporte_data, titulo="Reporte de Clientes"):
        """
        Genera PDF para reporte de clientes.
        """
        buffer = BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=A4, topMargin=0.75*inch, bottomMargin=0.75*inch)
        story = []
        
        # Título
        story.append(Paragraph(titulo, self.styles['TituloReporte']))
        story.append(Spacer(1, 0.2*inch))
        
        # Resumen
        resumen = reporte_data.get('resumen', {})
        fecha_generacion = f"<b>Fecha de generación:</b> {datetime.now().strftime('%d/%m/%Y %H:%M')}"
        story.append(Paragraph(fecha_generacion, self.styles['InfoReporte']))
        
        cantidad_clientes = reporte_data.get('cantidad_clientes', 0)
        story.append(Paragraph(f"<b>Total de clientes:</b> {cantidad_clientes}", self.styles['InfoReporte']))
        story.append(Spacer(1, 0.3*inch))
        
        # Tabla de clientes
        clientes = reporte_data.get('clientes', [])
        if clientes:
            story.append(Paragraph("Detalle de Clientes", self.styles['SubtituloReporte']))
            
            data_clientes = [[
                'Cliente', 'Email', 'Total Gastado', 
                'Reservas', 'Ticket Prom.', 'Última Compra'
            ]]
            
            for cliente in clientes[:50]:  # Limitar a 50 clientes por página
                ultima_compra = cliente.get('ultima_compra')
                if isinstance(ultima_compra, datetime):
                    ultima_compra = ultima_compra.strftime('%d/%m/%Y')
                
                data_clientes.append([
                    cliente.get('cliente__nombre', 'N/A')[:25],
                    cliente.get('cliente__user__email', 'N/A')[:30],  # Email del User de Django
                    self._formatear_valor(cliente.get('total_gastado')),
                    str(cliente.get('cantidad_reservas', 0)),
                    self._formatear_valor(cliente.get('ticket_promedio')),
                    str(ultima_compra) if ultima_compra else 'N/A'
                ])
            
            table_clientes = Table(data_clientes, colWidths=[
                1.5*inch, 1*inch, 1.2*inch, 0.8*inch, 1*inch, 1*inch
            ])
            table_clientes.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#3498DB')),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                ('ALIGN', (2, 1), (-1, -1), 'RIGHT'),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, 0), 8),
                ('FONTSIZE', (0, 1), (-1, -1), 7),
                ('BOTTOMPADDING', (0, 0), (-1, 0), 8),
                ('TOPPADDING', (0, 1), (-1, -1), 4),
                ('BOTTOMPADDING', (0, 1), (-1, -1), 4),
                ('GRID', (0, 0), (-1, -1), 0.5, colors.HexColor('#BDC3C7')),
                ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, colors.HexColor('#ECF0F1')])
            ]))
            story.append(table_clientes)
        
        doc.build(story)
        buffer.seek(0)
        return buffer
    
    def generar_reporte_productos(self, reporte_data, titulo="Reporte de Productos"):
        """
        Genera PDF para reporte de productos con filtros aplicados.
        """
        buffer = BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=letter, topMargin=0.75*inch, bottomMargin=0.75*inch)
        story = []
        filtros = reporte_data.get('filtros_aplicados', {})
        
        # Título
        story.append(Paragraph(titulo, self.styles['TituloReporte']))
        story.append(Spacer(1, 0.1*inch))
        
        # Fecha de generación
        fecha_generacion = f"<b>Fecha de generación:</b> {datetime.now().strftime('%d/%m/%Y %H:%M')}"
        story.append(Paragraph(fecha_generacion, self.styles['InfoReporte']))
        
        # Filtros aplicados
        filtros_texto = []
        if filtros.get('moneda'):
            filtros_texto.append(f"<b>Moneda:</b> {filtros['moneda']}")
        if filtros.get('fecha_inicio') or filtros.get('fecha_fin'):
            fecha_ini = filtros.get('fecha_inicio', 'N/A')
            fecha_fin = filtros.get('fecha_fin', 'N/A')
            if hasattr(fecha_ini, 'strftime'):
                fecha_ini = fecha_ini.strftime('%d/%m/%Y')
            if hasattr(fecha_fin, 'strftime'):
                fecha_fin = fecha_fin.strftime('%d/%m/%Y')
            filtros_texto.append(f"<b>Período:</b> {fecha_ini} - {fecha_fin}")
        if filtros.get('departamento'):
            filtros_texto.append(f"<b>Departamento:</b> {filtros['departamento']}")
        if filtros.get('ciudad'):
            filtros_texto.append(f"<b>Ciudad:</b> {filtros['ciudad']}")
        if filtros.get('tipo_destino'):
            filtros_texto.append(f"<b>Tipo Destino:</b> {filtros['tipo_destino']}")
        if filtros.get('tipo_cliente'):
            filtros_texto.append(f"<b>Tipo Cliente:</b> {filtros['tipo_cliente']}")
        if filtros.get('estado'):
            filtros_texto.append(f"<b>Estado:</b> {filtros['estado']}")
        
        # Mostrar filtros (máximo 4)
        for filtro_info in filtros_texto[:4]:
            story.append(Paragraph(filtro_info, self.styles['InfoReporte']))
        
        story.append(Spacer(1, 0.2*inch))
        
        # Paquetes
        paquetes = reporte_data.get('paquetes', [])
        moneda = filtros.get('moneda', 'USD').upper()
        
        if paquetes:
            story.append(Paragraph("Rendimiento de Paquetes", self.styles['SubtituloReporte']))
            
            # Encabezados según moneda seleccionada
            if moneda == 'BOB':
                # Solo BOB
                data_paquetes = [['Paquete', 'Precio Base (Bs)', 'Ventas Totales (Bs)', 'Cant.', 'Conv. %']]
                col_widths = [2.5*inch, 1.2*inch, 1.5*inch, 0.7*inch, 0.8*inch]
            elif moneda == 'USD':
                # Solo USD
                data_paquetes = [['Paquete', 'Precio Base (USD)', 'Ventas Totales (USD)', 'Cant.', 'Conv. %']]
                col_widths = [2.5*inch, 1.2*inch, 1.5*inch, 0.7*inch, 0.8*inch]
            else:
                # Ambas monedas (AMBAS o sin filtro)
                data_paquetes = [['Paquete', 'Precio USD', 'Precio BOB', 'Ventas BOB', 'Cant.', 'Conv. %']]
                col_widths = [2*inch, 0.9*inch, 1.1*inch, 1.2*inch, 0.6*inch, 0.7*inch]
            
            for paquete in paquetes[:20]:
                if moneda == 'BOB':
                    # Solo BOB
                    precio_base = paquete.get('paquete__precio_base_bob') or (paquete.get('paquete__precio_base', 0) * 6.96)
                    ventas_totales = paquete.get('ventas_totales_bob') or paquete.get('ventas_totales', 0)
                    data_paquetes.append([
                        paquete.get('paquete__nombre', 'N/A')[:30],
                        f"Bs. {float(precio_base):,.2f}",
                        f"Bs. {float(ventas_totales):,.2f}",
                        str(paquete.get('cantidad_vendida', 0)),
                        f"{float(paquete.get('tasa_conversion', 0)):.1f}%"
                    ])
                elif moneda == 'USD':
                    # Solo USD
                    precio_base = paquete.get('paquete__precio_base_usd') or paquete.get('paquete__precio_base', 0)
                    ventas_totales = paquete.get('ventas_totales_usd') or (paquete.get('ventas_totales', 0) / 6.96)
                    data_paquetes.append([
                        paquete.get('paquete__nombre', 'N/A')[:30],
                        f"${float(precio_base):,.2f}",
                        f"${float(ventas_totales):,.2f}",
                        str(paquete.get('cantidad_vendida', 0)),
                        f"{float(paquete.get('tasa_conversion', 0)):.1f}%"
                    ])
                else:
                    # Ambas monedas
                    data_paquetes.append([
                        paquete.get('paquete__nombre', 'N/A')[:25],
                        f"${float(paquete.get('paquete__precio_base_usd', 0)):,.2f}",
                        f"Bs. {float(paquete.get('paquete__precio_base_bob', 0)):,.2f}",
                        f"Bs. {float(paquete.get('ventas_totales_bob', 0)):,.2f}",
                        str(paquete.get('cantidad_vendida', 0)),
                        f"{float(paquete.get('tasa_conversion', 0)):.1f}%"
                    ])
            
            table_paquetes = Table(data_paquetes, colWidths=col_widths)
            table_paquetes.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#27AE60')),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                ('ALIGN', (1, 1), (-1, -1), 'RIGHT'),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, 0), 9),
                ('FONTSIZE', (0, 1), (-1, -1), 7.5),
                ('BOTTOMPADDING', (0, 0), (-1, 0), 10),
                ('TOPPADDING', (0, 1), (-1, -1), 5),
                ('BOTTOMPADDING', (0, 1), (-1, -1), 5),
                ('GRID', (0, 0), (-1, -1), 1, colors.HexColor('#BDC3C7')),
                ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, colors.HexColor('#ECF0F1')])
            ]))
            story.append(table_paquetes)
            story.append(Spacer(1, 0.3*inch))
        
        # Servicios
        servicios = reporte_data.get('servicios', [])
        if servicios:
            story.append(Paragraph("Rendimiento de Servicios", self.styles['SubtituloReporte']))
            
            # Encabezados según moneda seleccionada
            if moneda == 'BOB':
                # Solo BOB
                data_servicios = [['Servicio', 'Categoría', 'Precio (Bs)', 'Ventas Totales (Bs)', 'Cant.', 'Conv. %']]
                col_widths_serv = [2*inch, 1*inch, 1*inch, 1.5*inch, 0.6*inch, 0.7*inch]
            elif moneda == 'USD':
                # Solo USD
                data_servicios = [['Servicio', 'Categoría', 'Precio (USD)', 'Ventas Totales (USD)', 'Cant.', 'Conv. %']]
                col_widths_serv = [2*inch, 1*inch, 1*inch, 1.5*inch, 0.6*inch, 0.7*inch]
            else:
                # Ambas monedas
                data_servicios = [['Servicio', 'Cat.', 'USD', 'BOB', 'Ventas BOB', 'Cant.', 'Conv. %']]
                col_widths_serv = [1.5*inch, 0.7*inch, 0.7*inch, 0.8*inch, 1.2*inch, 0.5*inch, 0.6*inch]
            
            for servicio in servicios[:20]:
                if moneda == 'BOB':
                    # Solo BOB
                    precio_bob = servicio.get('servicio__precio_bob') or (servicio.get('servicio__precio_usd', 0) * 6.96)
                    ventas_totales = servicio.get('ventas_totales_bob') or servicio.get('ventas_totales', 0)
                    data_servicios.append([
                        servicio.get('servicio__titulo', 'N/A')[:30],
                        servicio.get('servicio__categoria__nombre', 'N/A')[:15],
                        f"Bs. {float(precio_bob):,.2f}",
                        f"Bs. {float(ventas_totales):,.2f}",
                        str(servicio.get('cantidad_vendida', 0)),
                        f"{float(servicio.get('tasa_conversion', 0)):.1f}%"
                    ])
                elif moneda == 'USD':
                    # Solo USD
                    precio_usd = servicio.get('servicio__precio_usd', 0)
                    ventas_totales = servicio.get('ventas_totales_usd') or (servicio.get('ventas_totales', 0) / 6.96)
                    data_servicios.append([
                        servicio.get('servicio__titulo', 'N/A')[:30],
                        servicio.get('servicio__categoria__nombre', 'N/A')[:15],
                        f"${float(precio_usd):,.2f}",
                        f"${float(ventas_totales):,.2f}",
                        str(servicio.get('cantidad_vendida', 0)),
                        f"{float(servicio.get('tasa_conversion', 0)):.1f}%"
                    ])
                else:
                    # Ambas monedas
                    data_servicios.append([
                        servicio.get('servicio__titulo', 'N/A')[:22],
                        servicio.get('servicio__categoria__nombre', 'N/A')[:12],
                        f"${float(servicio.get('servicio__precio_usd', 0)):,.2f}",
                        f"Bs. {float(servicio.get('servicio__precio_bob', 0)):,.0f}",
                        f"Bs. {float(servicio.get('ventas_totales_bob', 0)):,.2f}",
                        str(servicio.get('cantidad_vendida', 0)),
                        f"{float(servicio.get('tasa_conversion', 0)):.1f}%"
                    ])
            
            table_servicios = Table(data_servicios, colWidths=col_widths_serv)
            table_servicios.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#E67E22')),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                ('ALIGN', (2, 1), (-1, -1), 'RIGHT'),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, 0), 7.5),
                ('FONTSIZE', (0, 1), (-1, -1), 6.5),
                ('BOTTOMPADDING', (0, 0), (-1, 0), 8),
                ('TOPPADDING', (0, 1), (-1, -1), 4),
                ('BOTTOMPADDING', (0, 1), (-1, -1), 4),
                ('GRID', (0, 0), (-1, -1), 0.5, colors.HexColor('#BDC3C7')),
                ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, colors.HexColor('#ECF0F1')])
            ]))
            story.append(table_servicios)
        
        doc.build(story)
        buffer.seek(0)
        return buffer


class ExportadorReportesExcel:
    """
    Clase para generar reportes en formato Excel con estructura profesional.
    """
    
    def __init__(self, moneda='USD'):
        self.moneda = moneda  # USD o BOB
        self.simbolo_moneda = '$' if moneda == 'USD' else 'Bs.'
        self.tasa_cambio = 6.96  # 1 USD = 6.96 BOB
        self.color_header = 'FF3498DB'
        self.color_subtotal = 'FFECF0F1'
        self.border_style = Border(
            left=Side(style='thin', color='FFBDC3C7'),
            right=Side(style='thin', color='FFBDC3C7'),
            top=Side(style='thin', color='FFBDC3C7'),
            bottom=Side(style='thin', color='FFBDC3C7')
        )
    
    def _aplicar_estilo_header(self, ws, row_num, columns):
        """Aplica estilo a la fila de encabezado."""
        for col_num in range(1, columns + 1):
            cell = ws.cell(row=row_num, column=col_num)
            cell.fill = PatternFill(start_color=self.color_header, end_color=self.color_header, fill_type='solid')
            cell.font = Font(bold=True, color='FFFFFFFF', size=11)
            cell.alignment = Alignment(horizontal='center', vertical='center')
            cell.border = self.border_style
    
    def _aplicar_estilo_datos(self, ws, start_row, end_row, columns):
        """Aplica estilo a las filas de datos."""
        for row_num in range(start_row, end_row + 1):
            for col_num in range(1, columns + 1):
                cell = ws.cell(row=row_num, column=col_num)
                cell.border = self.border_style
                cell.alignment = Alignment(vertical='center')
                
                # Alternar colores de fila
                if row_num % 2 == 0:
                    cell.fill = PatternFill(start_color=self.color_subtotal, end_color=self.color_subtotal, fill_type='solid')
    
    def _ajustar_ancho_columnas(self, ws):
        """Ajusta automáticamente el ancho de las columnas."""
        for column in ws.columns:
            max_length = 0
            column_letter = get_column_letter(column[0].column)
            
            for cell in column:
                try:
                    if len(str(cell.value)) > max_length:
                        max_length = len(str(cell.value))
                except:
                    pass
            
            adjusted_width = min(max_length + 2, 50)
            ws.column_dimensions[column_letter].width = adjusted_width
    
    def generar_reporte_ventas_general(self, reporte_data, titulo="Reporte de Ventas General"):
        """
        Genera Excel para reporte de ventas general.
        """
        wb = Workbook()
        
        # Hoja 1: Resumen
        ws_resumen = cast(Worksheet, wb.active)
        ws_resumen.title = "Resumen"
        
        # Título
        ws_resumen['A1'] = titulo
        ws_resumen['A1'].font = Font(size=18, bold=True, color='FF2C3E50')
        ws_resumen.merge_cells('A1:D1')
        ws_resumen['A1'].alignment = Alignment(horizontal='center')
        
        # Información
        ws_resumen['A3'] = 'Fecha de generación:'
        ws_resumen['B3'] = datetime.now().strftime('%d/%m/%Y %H:%M')
        
        periodo = reporte_data.get('periodo', {})
        if periodo.get('fecha_inicio'):
            ws_resumen['A4'] = 'Período desde:'
            fecha_inicio = periodo['fecha_inicio']
            if isinstance(fecha_inicio, datetime):
                ws_resumen['B4'] = fecha_inicio.strftime('%d/%m/%Y')
            else:
                ws_resumen['B4'] = str(fecha_inicio)
        
        if periodo.get('fecha_fin'):
            ws_resumen['A5'] = 'Período hasta:'
            fecha_fin = periodo['fecha_fin']
            if isinstance(fecha_fin, datetime):
                ws_resumen['B5'] = fecha_fin.strftime('%d/%m/%Y')
            else:
                ws_resumen['B5'] = str(fecha_fin)
        
        # Métricas Generales
        row = 7
        ws_resumen[f'A{row}'] = 'MÉTRICAS GENERALES'
        ws_resumen[f'A{row}'].font = Font(size=14, bold=True, color='FF34495E')
        ws_resumen.merge_cells(f'A{row}:B{row}')
        
        row += 2
        metricas = reporte_data.get('metricas_generales', {})
        
        # Encabezados
        ws_resumen[f'A{row}'] = 'Métrica'
        ws_resumen[f'B{row}'] = 'Valor'
        self._aplicar_estilo_header(ws_resumen, row, 2)
        
        row += 1
        start_data_row = row
        
        metricas_lista = [
            ('Total Ventas (Bs)', metricas.get('total_ventas', 0)),
            ('Cantidad de Reservas', metricas.get('cantidad_reservas', 0)),
            ('Ticket Promedio (Bs)', metricas.get('ticket_promedio', 0)),
            ('Total Pagado (Bs)', metricas.get('total_pagado', 0)),
        ]
        
        for nombre, valor in metricas_lista:
            ws_resumen[f'A{row}'] = nombre
            if isinstance(valor, Decimal):
                ws_resumen[f'B{row}'] = float(valor)
                ws_resumen[f'B{row}'].number_format = '#,##0.00'
            else:
                ws_resumen[f'B{row}'] = valor
            row += 1
        
        self._aplicar_estilo_datos(ws_resumen, start_data_row, row - 1, 2)
        
        # Hoja 2: Top Paquetes
        top_paquetes = reporte_data.get('top_paquetes', [])
        if top_paquetes:
            ws_paquetes = cast(Worksheet, wb.create_sheet("Top Paquetes"))
            
            ws_paquetes['A1'] = 'Top Paquetes Más Vendidos'
            ws_paquetes['A1'].font = Font(size=16, bold=True, color='FF27AE60')
            ws_paquetes.merge_cells('A1:D1')
            
            # Encabezados
            headers = ['Paquete', 'ID', 'Total Ventas (Bs)', 'Cantidad']
            for col_num, header in enumerate(headers, 1):
                ws_paquetes.cell(row=3, column=col_num, value=header)
            self._aplicar_estilo_header(ws_paquetes, 3, len(headers))
            
            # Datos
            for row_num, paquete in enumerate(top_paquetes, 4):
                ws_paquetes.cell(row=row_num, column=1, value=paquete.get('paquete__nombre', 'N/A'))
                ws_paquetes.cell(row=row_num, column=2, value=paquete.get('paquete__id', ''))
                
                total_ventas = paquete.get('total_ventas')
                if total_ventas:
                    ws_paquetes.cell(row=row_num, column=3, value=float(total_ventas))
                    ws_paquetes.cell(row=row_num, column=3).number_format = '#,##0.00'
                
                ws_paquetes.cell(row=row_num, column=4, value=paquete.get('cantidad', 0))
            
            self._aplicar_estilo_datos(ws_paquetes, 4, 3 + len(top_paquetes), len(headers))
            self._ajustar_ancho_columnas(ws_paquetes)
        
        # Hoja 3: Top Servicios
        top_servicios = reporte_data.get('top_servicios', [])
        if top_servicios:
            ws_servicios = cast(Worksheet, wb.create_sheet("Top Servicios"))
            
            ws_servicios['A1'] = 'Top Servicios Más Vendidos'
            ws_servicios['A1'].font = Font(size=16, bold=True, color='FFE67E22')
            ws_servicios.merge_cells('A1:D1')
            
            # Encabezados
            headers = ['Servicio', 'ID', 'Total Ventas (Bs)', 'Cantidad']
            for col_num, header in enumerate(headers, 1):
                ws_servicios.cell(row=3, column=col_num, value=header)
            self._aplicar_estilo_header(ws_servicios, 3, len(headers))
            
            # Datos
            for row_num, servicio in enumerate(top_servicios, 4):
                ws_servicios.cell(row=row_num, column=1, value=servicio.get('servicio__titulo', 'N/A'))
                ws_servicios.cell(row=row_num, column=2, value=servicio.get('servicio__id', ''))
                
                total_ventas = servicio.get('total_ventas')
                if total_ventas:
                    ws_servicios.cell(row=row_num, column=3, value=float(total_ventas))
                    ws_servicios.cell(row=row_num, column=3).number_format = '#,##0.00'
                
                ws_servicios.cell(row=row_num, column=4, value=servicio.get('cantidad', 0))
            
            self._aplicar_estilo_datos(ws_servicios, 4, 3 + len(top_servicios), len(headers))
            self._ajustar_ancho_columnas(ws_servicios)
        
        # Hoja 4: Top Clientes
        top_clientes = reporte_data.get('top_clientes', [])
        if top_clientes:
            ws_clientes = cast(Worksheet, wb.create_sheet("Top Clientes"))
            
            ws_clientes['A1'] = 'Top Clientes'
            ws_clientes['A1'].font = Font(size=16, bold=True, color='FF9B59B6')
            ws_clientes.merge_cells('A1:D1')
            
            # Encabezados
            headers = ['Cliente', 'ID', 'Total Gastado (Bs)', 'Cantidad Reservas']
            for col_num, header in enumerate(headers, 1):
                ws_clientes.cell(row=3, column=col_num, value=header)
            self._aplicar_estilo_header(ws_clientes, 3, len(headers))
            
            # Datos
            for row_num, cliente in enumerate(top_clientes, 4):
                ws_clientes.cell(row=row_num, column=1, value=cliente.get('cliente__nombre', 'N/A'))
                ws_clientes.cell(row=row_num, column=2, value=cliente.get('cliente__id', ''))
                
                total_gastado = cliente.get('total_gastado')
                if total_gastado:
                    ws_clientes.cell(row=row_num, column=3, value=float(total_gastado))
                    ws_clientes.cell(row=row_num, column=3).number_format = '#,##0.00'
                
                ws_clientes.cell(row=row_num, column=4, value=cliente.get('cantidad_reservas', 0))
            
            self._aplicar_estilo_datos(ws_clientes, 4, 3 + len(top_clientes), len(headers))
            self._ajustar_ancho_columnas(ws_clientes)
        
        self._ajustar_ancho_columnas(ws_resumen)
        
        # Guardar en buffer
        buffer = BytesIO()
        wb.save(buffer)
        buffer.seek(0)
        return buffer
    
    def generar_reporte_clientes(self, reporte_data, titulo="Reporte de Clientes"):
        """
        Genera Excel para reporte de clientes.
        """
        wb = Workbook()
        ws = cast(Worksheet, wb.active)
        ws.title = "Clientes"
        
        # Título
        ws['A1'] = titulo
        ws['A1'].font = Font(size=18, bold=True, color='FF2C3E50')
        ws.merge_cells('A1:G1')
        ws['A1'].alignment = Alignment(horizontal='center')
        
        # Información
        ws['A3'] = f"Generado: {datetime.now().strftime('%d/%m/%Y %H:%M')}"
        ws['A4'] = f"Total clientes: {reporte_data.get('cantidad_clientes', 0)}"
        
        # Encabezados
        headers = ['Cliente', 'Email', 'Total Gastado (Bs)', 'Reservas', 
                   'Reservas Pagadas', 'Ticket Promedio (Bs)', 'Última Compra']
        for col_num, header in enumerate(headers, 1):
            ws.cell(row=6, column=col_num, value=header)
        self._aplicar_estilo_header(ws, 6, len(headers))
        
        # Datos
        clientes = reporte_data.get('clientes', [])
        for row_num, cliente in enumerate(clientes, 7):
            ws.cell(row=row_num, column=1, value=cliente.get('cliente__nombre', 'N/A'))
            ws.cell(row=row_num, column=2, value=cliente.get('cliente__user__email', 'N/A'))  # Email del User de Django
            
            total_gastado = cliente.get('total_gastado')
            if total_gastado:
                ws.cell(row=row_num, column=3, value=float(total_gastado))
                ws.cell(row=row_num, column=3).number_format = '#,##0.00'
            
            ws.cell(row=row_num, column=4, value=cliente.get('cantidad_reservas', 0))
            ws.cell(row=row_num, column=5, value=cliente.get('reservas_pagadas', 0))
            
            ticket_promedio = cliente.get('ticket_promedio')
            if ticket_promedio:
                ws.cell(row=row_num, column=6, value=float(ticket_promedio))
                ws.cell(row=row_num, column=6).number_format = '#,##0.00'
            
            ultima_compra = cliente.get('ultima_compra')
            if isinstance(ultima_compra, datetime):
                ws.cell(row=row_num, column=7, value=ultima_compra.strftime('%d/%m/%Y'))
            else:
                ws.cell(row=row_num, column=7, value=str(ultima_compra) if ultima_compra else 'N/A')
        
        if clientes:
            self._aplicar_estilo_datos(ws, 7, 6 + len(clientes), len(headers))
        
        self._ajustar_ancho_columnas(ws)
        
        buffer = BytesIO()
        wb.save(buffer)
        buffer.seek(0)
        return buffer
    
    def generar_reporte_productos(self, reporte_data, titulo="Reporte de Productos"):
        """
        Genera Excel para reporte de productos con filtros aplicados en la cabecera.
        """
        wb = Workbook()
        filtros = reporte_data.get('filtros_aplicados', {})
        
        # Hoja 1: Paquetes
        paquetes = reporte_data.get('paquetes', [])
        if paquetes:
            ws_paquetes = cast(Worksheet, wb.active)
            ws_paquetes.title = "Paquetes"
            
            # Título principal
            ws_paquetes['A1'] = 'Rendimiento de Paquetes'
            ws_paquetes['A1'].font = Font(size=16, bold=True, color='FF27AE60')
            ws_paquetes.merge_cells('A1:F1')
            ws_paquetes['A1'].alignment = Alignment(horizontal='center')
            
            # Información de filtros (fila 2)
            row_info = 2
            ws_paquetes[f'A{row_info}'] = f"Generado: {datetime.now().strftime('%d/%m/%Y %H:%M')}"
            ws_paquetes[f'A{row_info}'].font = Font(size=9, color='FF7F8C8D')
            
            # Mostrar filtros aplicados (fila 3-5)
            filtros_texto = []
            if filtros.get('moneda'):
                filtros_texto.append(f"Moneda: {filtros['moneda']}")
            if filtros.get('fecha_inicio') or filtros.get('fecha_fin'):
                fecha_ini = filtros.get('fecha_inicio', 'N/A')
                fecha_fin = filtros.get('fecha_fin', 'N/A')
                if hasattr(fecha_ini, 'strftime'):
                    fecha_ini = fecha_ini.strftime('%d/%m/%Y')
                if hasattr(fecha_fin, 'strftime'):
                    fecha_fin = fecha_fin.strftime('%d/%m/%Y')
                filtros_texto.append(f"Período: {fecha_ini} - {fecha_fin}")
            if filtros.get('departamento'):
                filtros_texto.append(f"Departamento: {filtros['departamento']}")
            if filtros.get('ciudad'):
                filtros_texto.append(f"Ciudad: {filtros['ciudad']}")
            if filtros.get('tipo_destino'):
                filtros_texto.append(f"Tipo Destino: {filtros['tipo_destino']}")
            if filtros.get('tipo_cliente'):
                filtros_texto.append(f"Tipo Cliente: {filtros['tipo_cliente']}")
            if filtros.get('estado'):
                filtros_texto.append(f"Estado: {filtros['estado']}")
            
            # Escribir filtros en filas 3, 4, 5
            for idx, filtro_info in enumerate(filtros_texto[:3], start=3):
                ws_paquetes[f'A{idx}'] = filtro_info
                ws_paquetes[f'A{idx}'].font = Font(size=9, bold=True, color='FF34495E')
            
            # Total de registros
            ws_paquetes['E3'] = f"Total paquetes: {len(paquetes)}"
            ws_paquetes['E3'].font = Font(size=9, bold=True, color='FF27AE60')
            ws_paquetes['E3'].alignment = Alignment(horizontal='right')
            
            # Obtener moneda seleccionada
            moneda = filtros.get('moneda', 'USD').upper()
            
            # Encabezados de tabla según moneda (fila 6)
            if moneda == 'BOB':
                headers = ['Paquete', 'Precio Base (Bs)', 'Ventas Totales (Bs)', 
                          'Cantidad Vendida', 'Tasa Conversión (%)', '¿Personalizado?']
            elif moneda == 'USD':
                headers = ['Paquete', 'Precio Base (USD)', 'Ventas Totales (USD)', 
                          'Cantidad Vendida', 'Tasa Conversión (%)', '¿Personalizado?']
            else:
                # Ambas monedas
                headers = ['Paquete', 'Precio USD', 'Precio BOB', 'Ventas BOB', 
                          'Cantidad Vendida', 'Tasa Conversión (%)', '¿Personalizado?']
            
            for col_num, header in enumerate(headers, 1):
                ws_paquetes.cell(row=6, column=col_num, value=header)
            self._aplicar_estilo_header(ws_paquetes, 6, len(headers))
            
            # Datos de paquetes (desde fila 7)
            for row_num, paquete in enumerate(paquetes, 7):
                ws_paquetes.cell(row=row_num, column=1, value=paquete.get('paquete__nombre', 'N/A'))
                
                if moneda == 'BOB':
                    # Solo BOB
                    precio_base = paquete.get('paquete__precio_base_bob') or (paquete.get('paquete__precio_base', 0) * 6.96)
                    if precio_base:
                        ws_paquetes.cell(row=row_num, column=2, value=float(precio_base))
                        ws_paquetes.cell(row=row_num, column=2).number_format = '#,##0.00'
                    
                    ventas_totales = paquete.get('ventas_totales_bob') or paquete.get('ventas_totales', 0)
                    if ventas_totales:
                        ws_paquetes.cell(row=row_num, column=3, value=float(ventas_totales))
                        ws_paquetes.cell(row=row_num, column=3).number_format = '#,##0.00'
                    
                    ws_paquetes.cell(row=row_num, column=4, value=paquete.get('cantidad_vendida', 0))
                    
                    tasa_conversion = paquete.get('tasa_conversion', 0)
                    ws_paquetes.cell(row=row_num, column=5, value=float(tasa_conversion))
                    ws_paquetes.cell(row=row_num, column=5).number_format = '0.00'
                    
                    es_personalizado = paquete.get('paquete__es_personalizado', False)
                    ws_paquetes.cell(row=row_num, column=6, value='Sí' if es_personalizado else 'No')
                
                elif moneda == 'USD':
                    # Solo USD
                    precio_base = paquete.get('paquete__precio_base_usd') or paquete.get('paquete__precio_base', 0)
                    if precio_base:
                        ws_paquetes.cell(row=row_num, column=2, value=float(precio_base))
                        ws_paquetes.cell(row=row_num, column=2).number_format = '$#,##0.00'
                    
                    ventas_totales = paquete.get('ventas_totales_usd') or (paquete.get('ventas_totales', 0) / 6.96)
                    if ventas_totales:
                        ws_paquetes.cell(row=row_num, column=3, value=float(ventas_totales))
                        ws_paquetes.cell(row=row_num, column=3).number_format = '$#,##0.00'
                    
                    ws_paquetes.cell(row=row_num, column=4, value=paquete.get('cantidad_vendida', 0))
                    
                    tasa_conversion = paquete.get('tasa_conversion', 0)
                    ws_paquetes.cell(row=row_num, column=5, value=float(tasa_conversion))
                    ws_paquetes.cell(row=row_num, column=5).number_format = '0.00'
                    
                    es_personalizado = paquete.get('paquete__es_personalizado', False)
                    ws_paquetes.cell(row=row_num, column=6, value='Sí' if es_personalizado else 'No')
                
                else:
                    # Ambas monedas
                    precio_usd = paquete.get('paquete__precio_base_usd') or paquete.get('paquete__precio_base', 0)
                    precio_bob = paquete.get('paquete__precio_base_bob') or (precio_usd * 6.96)
                    
                    ws_paquetes.cell(row=row_num, column=2, value=float(precio_usd))
                    ws_paquetes.cell(row=row_num, column=2).number_format = '$#,##0.00'
                    
                    ws_paquetes.cell(row=row_num, column=3, value=float(precio_bob))
                    ws_paquetes.cell(row=row_num, column=3).number_format = '#,##0.00'
                    
                    ventas_totales = paquete.get('ventas_totales_bob') or paquete.get('ventas_totales', 0)
                    ws_paquetes.cell(row=row_num, column=4, value=float(ventas_totales))
                    ws_paquetes.cell(row=row_num, column=4).number_format = '#,##0.00'
                    
                    ws_paquetes.cell(row=row_num, column=5, value=paquete.get('cantidad_vendida', 0))
                    
                    tasa_conversion = paquete.get('tasa_conversion', 0)
                    ws_paquetes.cell(row=row_num, column=6, value=float(tasa_conversion))
                    ws_paquetes.cell(row=row_num, column=6).number_format = '0.00'
                    
                    es_personalizado = paquete.get('paquete__es_personalizado', False)
                    ws_paquetes.cell(row=row_num, column=7, value='Sí' if es_personalizado else 'No')
            
            self._aplicar_estilo_datos(ws_paquetes, 7, 6 + len(paquetes), len(headers))
            self._ajustar_ancho_columnas(ws_paquetes)
        
        # Hoja 2: Servicios
        servicios = reporte_data.get('servicios', [])
        if servicios:
            ws_servicios = cast(Worksheet, wb.create_sheet("Servicios"))
            
            # Título principal
            ws_servicios['A1'] = 'Rendimiento de Servicios'
            ws_servicios['A1'].font = Font(size=16, bold=True, color='FFE67E22')
            ws_servicios.merge_cells('A1:F1')
            ws_servicios['A1'].alignment = Alignment(horizontal='center')
            
            # Información de filtros (similar a Paquetes)
            row_info = 2
            ws_servicios[f'A{row_info}'] = f"Generado: {datetime.now().strftime('%d/%m/%Y %H:%M')}"
            ws_servicios[f'A{row_info}'].font = Font(size=9, color='FF7F8C8D')
            
            # Escribir filtros (reutilizar filtros_texto si existe)
            filtros_texto = []
            if filtros.get('moneda'):
                filtros_texto.append(f"Moneda: {filtros['moneda']}")
            if filtros.get('fecha_inicio') or filtros.get('fecha_fin'):
                fecha_ini = filtros.get('fecha_inicio', 'N/A')
                fecha_fin = filtros.get('fecha_fin', 'N/A')
                if hasattr(fecha_ini, 'strftime'):
                    fecha_ini = fecha_ini.strftime('%d/%m/%Y')
                if hasattr(fecha_fin, 'strftime'):
                    fecha_fin = fecha_fin.strftime('%d/%m/%Y')
                filtros_texto.append(f"Período: {fecha_ini} - {fecha_fin}")
            if filtros.get('categoria'):
                filtros_texto.append(f"Categoría: {filtros['categoria']}")
            
            for idx, filtro_info in enumerate(filtros_texto[:3], start=3):
                ws_servicios[f'A{idx}'] = filtro_info
                ws_servicios[f'A{idx}'].font = Font(size=9, bold=True, color='FF34495E')
            
            # Total de registros
            ws_servicios['E3'] = f"Total servicios: {len(servicios)}"
            ws_servicios['E3'].font = Font(size=9, bold=True, color='FFE67E22')
            ws_servicios['E3'].alignment = Alignment(horizontal='right')
            
            # Obtener moneda seleccionada
            moneda = filtros.get('moneda', 'USD').upper()
            
            # Encabezados de tabla según moneda (fila 6)
            if moneda == 'BOB':
                headers = ['Servicio', 'Categoría', 'Precio (Bs)', 
                          'Ventas Totales (Bs)', 'Cantidad Vendida', 'Tasa Conversión (%)']
            elif moneda == 'USD':
                headers = ['Servicio', 'Categoría', 'Precio (USD)', 
                          'Ventas Totales (USD)', 'Cantidad Vendida', 'Tasa Conversión (%)']
            else:
                # Ambas monedas
                headers = ['Servicio', 'Categoría', 'Precio USD', 'Precio BOB',
                          'Ventas BOB', 'Cantidad Vendida', 'Tasa Conversión (%)']
            
            for col_num, header in enumerate(headers, 1):
                ws_servicios.cell(row=6, column=col_num, value=header)
            self._aplicar_estilo_header(ws_servicios, 6, len(headers))
            
            # Datos de servicios (desde fila 7)
            for row_num, servicio in enumerate(servicios, 7):
                ws_servicios.cell(row=row_num, column=1, value=servicio.get('servicio__titulo', 'N/A'))
                ws_servicios.cell(row=row_num, column=2, value=servicio.get('servicio__categoria__nombre', 'N/A'))
                
                if moneda == 'BOB':
                    # Solo BOB
                    precio_bob = servicio.get('servicio__precio_bob') or (servicio.get('servicio__precio_usd', 0) * 6.96)
                    ws_servicios.cell(row=row_num, column=3, value=float(precio_bob))
                    ws_servicios.cell(row=row_num, column=3).number_format = '#,##0.00'
                    
                    ventas_totales = servicio.get('ventas_totales_bob') or servicio.get('ventas_totales', 0)
                    ws_servicios.cell(row=row_num, column=4, value=float(ventas_totales))
                    ws_servicios.cell(row=row_num, column=4).number_format = '#,##0.00'
                    
                    ws_servicios.cell(row=row_num, column=5, value=servicio.get('cantidad_vendida', 0))
                    
                    tasa_conversion = servicio.get('tasa_conversion', 0)
                    ws_servicios.cell(row=row_num, column=6, value=float(tasa_conversion))
                    ws_servicios.cell(row=row_num, column=6).number_format = '0.00'
                
                elif moneda == 'USD':
                    # Solo USD
                    precio_usd = servicio.get('servicio__precio_usd', 0)
                    ws_servicios.cell(row=row_num, column=3, value=float(precio_usd))
                    ws_servicios.cell(row=row_num, column=3).number_format = '$#,##0.00'
                    
                    ventas_totales = servicio.get('ventas_totales_usd') or (servicio.get('ventas_totales', 0) / 6.96)
                    ws_servicios.cell(row=row_num, column=4, value=float(ventas_totales))
                    ws_servicios.cell(row=row_num, column=4).number_format = '$#,##0.00'
                    
                    ws_servicios.cell(row=row_num, column=5, value=servicio.get('cantidad_vendida', 0))
                    
                    tasa_conversion = servicio.get('tasa_conversion', 0)
                    ws_servicios.cell(row=row_num, column=6, value=float(tasa_conversion))
                    ws_servicios.cell(row=row_num, column=6).number_format = '0.00'
                
                else:
                    # Ambas monedas
                    precio_usd = servicio.get('servicio__precio_usd', 0)
                    precio_bob = servicio.get('servicio__precio_bob') or (precio_usd * 6.96)
                    
                    ws_servicios.cell(row=row_num, column=3, value=float(precio_usd))
                    ws_servicios.cell(row=row_num, column=3).number_format = '$#,##0.00'
                    
                    ws_servicios.cell(row=row_num, column=4, value=float(precio_bob))
                    ws_servicios.cell(row=row_num, column=4).number_format = '#,##0.00'
                    
                    ventas_totales = servicio.get('ventas_totales_bob') or servicio.get('ventas_totales', 0)
                    ws_servicios.cell(row=row_num, column=5, value=float(ventas_totales))
                    ws_servicios.cell(row=row_num, column=5).number_format = '#,##0.00'
                    
                    ws_servicios.cell(row=row_num, column=6, value=servicio.get('cantidad_vendida', 0))
                    
                    tasa_conversion = servicio.get('tasa_conversion', 0)
                    ws_servicios.cell(row=row_num, column=7, value=float(tasa_conversion))
                    ws_servicios.cell(row=row_num, column=7).number_format = '0.00'
            
            self._aplicar_estilo_datos(ws_servicios, 7, 6 + len(servicios), len(headers))
            self._ajustar_ancho_columnas(ws_servicios)
        
        buffer = BytesIO()
        wb.save(buffer)
        buffer.seek(0)
        return buffer


class ExportadorReportesWord:
    """
    Clase para generar reportes de clientes en formato Word (DOCX) con estructura profesional.
    """
    
    def __init__(self):
        """Inicializa el exportador de Word."""
        self.colores = {
            'titulo': RGBColor(44, 62, 80),  # #2C3E50
            'subtitulo': RGBColor(52, 73, 94),  # #34495E
            'header_table': RGBColor(41, 128, 185),  # #2980B9
            'texto': RGBColor(0, 0, 0),
        }
    
    def generar_reporte_clientes(self, datos_reporte):
        """
        Genera un reporte de clientes en formato Word.
        
        Args:
            datos_reporte (dict): Diccionario con los datos del reporte
                - resumen (dict): Resumen general
                - clientes (list): Lista de clientes
                - filtros (dict): Filtros aplicados
        
        Returns:
            BytesIO: Buffer con el documento Word generado
        """
        doc = Document()
        
        # Configurar márgenes
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(0.75)
            section.bottom_margin = Inches(0.75)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)
        
        resumen = datos_reporte.get('resumen', {})
        clientes = datos_reporte.get('clientes', [])
        filtros = datos_reporte.get('filtros', {})
        
        # ============================================
        # 1. TÍTULO PRINCIPAL
        # ============================================
        titulo = doc.add_heading('Reporte Detallado de Clientes', level=0)
        titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
        titulo_run = titulo.runs[0]
        titulo_run.font.color.rgb = self.colores['titulo']
        titulo_run.font.size = Pt(24)
        titulo_run.font.bold = True
        
        # Fecha de generación
        fecha_generacion = doc.add_paragraph()
        fecha_generacion.alignment = WD_ALIGN_PARAGRAPH.CENTER
        fecha_run = fecha_generacion.add_run(
            f'Fecha de generación: {datetime.now().strftime("%d/%m/%Y %H:%M:%S")}'
        )
        fecha_run.font.size = Pt(10)
        fecha_run.italic = True
        
        doc.add_paragraph()  # Espacio
        
        # ============================================
        # 2. INFORMACIÓN DE FILTROS APLICADOS
        # ============================================
        if filtros:
            filtros_heading = doc.add_heading('Filtros Aplicados', level=1)
            filtros_heading.runs[0].font.color.rgb = self.colores['subtitulo']
            
            # Moneda
            moneda = filtros.get('moneda', 'USD')
            p_moneda = doc.add_paragraph()
            p_moneda.add_run('Moneda: ').bold = True
            p_moneda.add_run(moneda)
            
            # Rango de fechas
            if filtros.get('fecha_inicio') or filtros.get('fecha_fin'):
                fecha_inicio = filtros.get('fecha_inicio', 'N/A')
                fecha_fin = filtros.get('fecha_fin', 'N/A')
                
                if isinstance(fecha_inicio, datetime):
                    fecha_inicio = fecha_inicio.strftime('%d/%m/%Y')
                if isinstance(fecha_fin, datetime):
                    fecha_fin = fecha_fin.strftime('%d/%m/%Y')
                
                p_periodo = doc.add_paragraph()
                p_periodo.add_run('Período: ').bold = True
                p_periodo.add_run(f'{fecha_inicio} - {fecha_fin}')
            
            # Tipo de cliente
            if filtros.get('tipo_cliente'):
                p_tipo = doc.add_paragraph()
                p_tipo.add_run('Tipo de Cliente: ').bold = True
                p_tipo.add_run(filtros['tipo_cliente'].title())
            
            # Rangos de monto
            if filtros.get('monto_minimo') or filtros.get('monto_maximo'):
                monto_min = filtros.get('monto_minimo', 0)
                monto_max = filtros.get('monto_maximo', '∞')
                
                p_monto = doc.add_paragraph()
                p_monto.add_run('Rango de Monto: ').bold = True
                simbolo_moneda = '$' if moneda == 'USD' else 'Bs'
                p_monto.add_run(f'{simbolo_moneda}{monto_min} - {simbolo_moneda}{monto_max}')
            
            doc.add_paragraph()  # Espacio
        
        # ============================================
        # 3. RESUMEN GENERAL
        # ============================================
        resumen_heading = doc.add_heading('Resumen General', level=1)
        resumen_heading.runs[0].font.color.rgb = self.colores['subtitulo']
        
        # Tabla de resumen (2 columnas)
        table_resumen = doc.add_table(rows=4, cols=2)
        table_resumen.style = 'Light Grid Accent 1'
        table_resumen.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        # Encabezados de resumen
        headers_resumen = table_resumen.rows[0].cells
        headers_resumen[0].text = 'Métrica'
        headers_resumen[1].text = 'Valor'
        
        for cell in headers_resumen:
            cell_paragraph = cell.paragraphs[0]
            cell_run = cell_paragraph.runs[0]
            cell_run.font.bold = True
            cell_run.font.color.rgb = RGBColor(255, 255, 255)
            shading_elm = cell._element.get_or_add_tcPr().get_or_add_shd()
            shading_elm.set_val("clear")
            shading_elm.set_fill("2980B9")  # Azul #2980B9
            cell_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Datos de resumen
        moneda = filtros.get('moneda', 'USD')
        simbolo = '$' if moneda == 'USD' else 'Bs'
        
        datos_resumen = [
            ['Total de Clientes', f"{resumen.get('total_clientes', 0)}"],
            ['Total Reservas', f"{resumen.get('total_reservas', 0)}"],
            ['Ingresos Totales', f"{simbolo}{resumen.get('ingresos_totales', 0):,.2f}"],
            ['Promedio por Cliente', f"{simbolo}{resumen.get('promedio_por_cliente', 0):,.2f}"],
        ]
        
        for i, (label, value) in enumerate(datos_resumen, start=1):
            row_cells = table_resumen.rows[i].cells
            row_cells[0].text = label
            row_cells[1].text = value
            
            # Negrita en etiquetas
            row_cells[0].paragraphs[0].runs[0].font.bold = True
            # Centrar valores
            row_cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph()  # Espacio
        
        # ============================================
        # 4. DETALLE DE CLIENTES
        # ============================================
        clientes_heading = doc.add_heading('Detalle de Clientes', level=1)
        clientes_heading.runs[0].font.color.rgb = self.colores['subtitulo']
        
        if not clientes:
            doc.add_paragraph('No se encontraron clientes con los filtros aplicados.')
        else:
            # Tabla de clientes
            num_cols = 5
            table_clientes = doc.add_table(rows=1, cols=num_cols)
            table_clientes.style = 'Light Grid Accent 1'
            
            # Encabezados
            headers = ['Cliente', 'Email', 'Total Reservas', 
                      f'Total Gastado ({simbolo})', 'Tipo Cliente']
            
            header_cells = table_clientes.rows[0].cells
            for i, header in enumerate(headers):
                header_cells[i].text = header
                cell_paragraph = header_cells[i].paragraphs[0]
                cell_run = cell_paragraph.runs[0]
                cell_run.font.bold = True
                cell_run.font.color.rgb = RGBColor(255, 255, 255)
                cell_run.font.size = Pt(10)
                
                # Fondo azul para encabezados
                shading_elm = header_cells[i]._element.get_or_add_tcPr().get_or_add_shd()
                shading_elm.set_val("clear")
                shading_elm.set_fill("2980B9")  # Azul #2980B9
                
                cell_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Datos de clientes
            for cliente in clientes:
                row_cells = table_clientes.add_row().cells
                
                # Nombre del cliente
                nombre_cliente = cliente.get('cliente__user__first_name', '') + ' ' + cliente.get('cliente__user__last_name', '')
                row_cells[0].text = nombre_cliente.strip() or 'N/A'
                
                # Email
                row_cells[1].text = cliente.get('cliente__user__email', 'N/A')
                
                # Total reservas
                row_cells[2].text = str(cliente.get('total_reservas', 0))
                row_cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                # Total gastado
                if moneda == 'BOB':
                    total = cliente.get('total_gastado_bob', 0) or cliente.get('total_gastado', 0)
                else:
                    total = cliente.get('total_gastado_usd', 0) or (cliente.get('total_gastado', 0) / 6.96)
                
                row_cells[3].text = f"{simbolo}{float(total):,.2f}"
                row_cells[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
                
                # Tipo de cliente
                tipo_cliente = cliente.get('tipo_cliente', 'N/A')
                row_cells[4].text = tipo_cliente.title()
                row_cells[4].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                # Ajustar tamaño de fuente en celdas
                for cell in row_cells:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.size = Pt(9)
            
            # Ajustar anchos de columna
            for i, width in enumerate([Inches(1.8), Inches(2.2), Inches(1.0), Inches(1.5), Inches(1.2)]):
                for row in table_clientes.rows:
                    row.cells[i].width = width
        
        # ============================================
        # 5. PIE DE PÁGINA
        # ============================================
        doc.add_paragraph()
        footer = doc.add_paragraph()
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer_run = footer.add_run('_' * 80)
        footer_run.font.color.rgb = RGBColor(150, 150, 150)
        
        info_footer = doc.add_paragraph()
        info_footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        info_run = info_footer.add_run(
            f'Sistema de Reservas Turísticas | Generado el {datetime.now().strftime("%d/%m/%Y a las %H:%M")}'
        )
        info_run.font.size = Pt(8)
        info_run.font.color.rgb = RGBColor(150, 150, 150)
        info_run.italic = True
        
        # Guardar en buffer
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer
    
    def generar_reporte_ventas(self, datos_reporte):
        """
        Genera un reporte de ventas en formato Word.
        
        Args:
            datos_reporte (dict): Diccionario con los datos del reporte
                - resumen (dict): Resumen general de ventas
                - ventas (list): Lista de ventas
                - filtros (dict): Filtros aplicados
        
        Returns:
            BytesIO: Buffer con el documento Word generado
        """
        doc = Document()
        
        # Configurar márgenes
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(0.75)
            section.bottom_margin = Inches(0.75)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)
        
        resumen = datos_reporte.get('resumen', {})
        ventas = datos_reporte.get('ventas', [])
        filtros = datos_reporte.get('filtros', {})
        
        # ============================================
        # 1. TÍTULO PRINCIPAL
        # ============================================
        titulo = doc.add_heading('Reporte de Ventas', level=1)
        titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in titulo.runs:
            run.font.color.rgb = self.colores['titulo']
            run.font.size = Pt(24)
            run.font.bold = True
        
        # Subtítulo con fecha
        subtitulo = doc.add_paragraph()
        subtitulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
        sub_run = subtitulo.add_run(f'Generado el {datetime.now().strftime("%d de %B de %Y")}')
        sub_run.font.size = Pt(12)
        sub_run.font.color.rgb = self.colores['subtitulo']
        sub_run.italic = True
        
        doc.add_paragraph()  # Espacio
        
        # ============================================
        # 2. FILTROS APLICADOS
        # ============================================
        if filtros:
            filtros_heading = doc.add_heading('Filtros Aplicados', level=2)
            for run in filtros_heading.runs:
                run.font.color.rgb = self.colores['subtitulo']
            
            filtros_text = []
            if filtros.get('fecha_inicio') and filtros.get('fecha_fin'):
                filtros_text.append(f"Período: {filtros['fecha_inicio']} - {filtros['fecha_fin']}")
            if filtros.get('monto_minimo'):
                filtros_text.append(f"Monto mínimo: ${filtros['monto_minimo']}")
            if filtros.get('monto_maximo'):
                filtros_text.append(f"Monto máximo: ${filtros['monto_maximo']}")
            if filtros.get('moneda'):
                filtros_text.append(f"Moneda: {filtros['moneda']}")
            
            filtros_p = doc.add_paragraph(' | '.join(filtros_text))
            filtros_p.runs[0].font.size = Pt(10)
            
            doc.add_paragraph()  # Espacio
        
        # ============================================
        # 3. RESUMEN EN TABLA
        # ============================================
        resumen_heading = doc.add_heading('Resumen General', level=2)
        for run in resumen_heading.runs:
            run.font.color.rgb = self.colores['subtitulo']
        
        # Tabla de resumen (2 columnas x 4 filas)
        table_resumen = doc.add_table(rows=4, cols=2)
        table_resumen.style = 'Light Grid Accent 1'
        
        # Encabezados
        hdr_cells = table_resumen.rows[0].cells
        hdr_cells[0].text = 'Métrica'
        hdr_cells[1].text = 'Valor'
        
        # Formato de encabezados
        for cell in hdr_cells:
            shading_elm = parse_xml(r'<w:shd {} w:fill="2980B9"/>'.format(nsdecls('w')))
            cell._tc.get_or_add_tcPr().append(shading_elm)
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
                    run.font.color.rgb = RGBColor(255, 255, 255)
                    run.font.size = Pt(11)
        
        # Datos de resumen
        row_data = [
            ('Total de Ventas', f"${resumen.get('total_ventas', 0):,.2f}"),
            ('Número de Transacciones', str(resumen.get('total_transacciones', 0))),
            ('Ticket Promedio', f"${resumen.get('ticket_promedio', 0):,.2f}"),
        ]
        
        for i, (metrica, valor) in enumerate(row_data, start=1):
            cells = table_resumen.rows[i].cells
            cells[0].text = metrica
            cells[1].text = valor
            
            # Formato de celdas
            for cell in cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(10)
        
        doc.add_paragraph()  # Espacio
        
        # ============================================
        # 4. DETALLE DE VENTAS
        # ============================================
        if ventas:
            detalle_heading = doc.add_heading('Detalle de Ventas', level=2)
            for run in detalle_heading.runs:
                run.font.color.rgb = self.colores['subtitulo']
            
            # Tabla de ventas
            table_ventas = doc.add_table(rows=1, cols=5)
            table_ventas.style = 'Light Grid Accent 1'
            
            # Encabezados
            hdr_cells = table_ventas.rows[0].cells
            headers = ['Cliente', 'Producto', 'Fecha', 'Monto', 'Estado']
            for i, header in enumerate(headers):
                hdr_cells[i].text = header
            
            # Formato de encabezados
            for cell in hdr_cells:
                shading_elm = parse_xml(r'<w:shd {} w:fill="2980B9"/>'.format(nsdecls('w')))
                cell._tc.get_or_add_tcPr().append(shading_elm)
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True
                        run.font.color.rgb = RGBColor(255, 255, 255)
                        run.font.size = Pt(11)
            
            # Datos de ventas
            for venta in ventas[:50]:  # Limitar a 50 ventas
                row_cells = table_ventas.add_row().cells
                row_cells[0].text = venta.get('cliente', 'N/A')
                row_cells[1].text = venta.get('producto', 'N/A')
                row_cells[2].text = venta.get('fecha', 'N/A')
                row_cells[3].text = f"${venta.get('monto', 0):,.2f}"
                row_cells[4].text = venta.get('estado', 'N/A')
                
                # Ajustar tamaño de fuente
                for cell in row_cells:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.size = Pt(9)
            
            # Ajustar anchos de columna
            for i, width in enumerate([Inches(1.8), Inches(2.2), Inches(1.0), Inches(1.2), Inches(1.0)]):
                for row in table_ventas.rows:
                    row.cells[i].width = width
        
        # ============================================
        # 5. PIE DE PÁGINA
        # ============================================
        doc.add_paragraph()
        footer = doc.add_paragraph()
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer_run = footer.add_run('_' * 80)
        footer_run.font.color.rgb = RGBColor(150, 150, 150)
        
        info_footer = doc.add_paragraph()
        info_footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        info_run = info_footer.add_run(
            f'Sistema de Reservas Turísticas | Generado el {datetime.now().strftime("%d/%m/%Y a las %H:%M")}'
        )
        info_run.font.size = Pt(8)
        info_run.font.color.rgb = RGBColor(150, 150, 150)
        info_run.italic = True
        
        # Guardar en buffer
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer
    
    def generar_reporte_productos(self, datos_reporte):
        """
        Genera un reporte de productos/paquetes en formato Word.
        
        Args:
            datos_reporte (dict): Diccionario con los datos del reporte
                - resumen (dict): Resumen general
                - productos (list): Lista de productos
                - filtros (dict): Filtros aplicados
        
        Returns:
            BytesIO: Buffer con el documento Word generado
        """
        doc = Document()
        
        # Configurar márgenes
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(0.75)
            section.bottom_margin = Inches(0.75)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)
        
        resumen = datos_reporte.get('resumen', {})
        productos = datos_reporte.get('productos', [])
        filtros = datos_reporte.get('filtros', {})
        
        # ============================================
        # 1. TÍTULO PRINCIPAL
        # ============================================
        titulo = doc.add_heading('Reporte de Productos y Paquetes', level=1)
        titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in titulo.runs:
            run.font.color.rgb = self.colores['titulo']
            run.font.size = Pt(24)
            run.font.bold = True
        
        # Subtítulo con fecha
        subtitulo = doc.add_paragraph()
        subtitulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
        sub_run = subtitulo.add_run(f'Generado el {datetime.now().strftime("%d de %B de %Y")}')
        sub_run.font.size = Pt(12)
        sub_run.font.color.rgb = self.colores['subtitulo']
        sub_run.italic = True
        
        doc.add_paragraph()  # Espacio
        
        # ============================================
        # 2. FILTROS APLICADOS
        # ============================================
        if filtros:
            filtros_heading = doc.add_heading('Filtros Aplicados', level=2)
            for run in filtros_heading.runs:
                run.font.color.rgb = self.colores['subtitulo']
            
            filtros_text = []
            if filtros.get('departamento'):
                filtros_text.append(f"Departamento: {filtros['departamento']}")
            if filtros.get('moneda'):
                filtros_text.append(f"Moneda: {filtros['moneda']}")
            if filtros.get('tipo_producto'):
                filtros_text.append(f"Tipo: {filtros['tipo_producto']}")
            
            if filtros_text:
                filtros_p = doc.add_paragraph(' | '.join(filtros_text))
                filtros_p.runs[0].font.size = Pt(10)
            
            doc.add_paragraph()  # Espacio
        
        # ============================================
        # 3. RESUMEN EN TABLA
        # ============================================
        resumen_heading = doc.add_heading('Resumen General', level=2)
        for run in resumen_heading.runs:
            run.font.color.rgb = self.colores['subtitulo']
        
        # Tabla de resumen (2 columnas x 4 filas)
        table_resumen = doc.add_table(rows=4, cols=2)
        table_resumen.style = 'Light Grid Accent 1'
        
        # Encabezados
        hdr_cells = table_resumen.rows[0].cells
        hdr_cells[0].text = 'Métrica'
        hdr_cells[1].text = 'Valor'
        
        # Formato de encabezados
        for cell in hdr_cells:
            shading_elm = parse_xml(r'<w:shd {} w:fill="2980B9"/>'.format(nsdecls('w')))
            cell._tc.get_or_add_tcPr().append(shading_elm)
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
                    run.font.color.rgb = RGBColor(255, 255, 255)
                    run.font.size = Pt(11)
        
        # Datos de resumen
        row_data = [
            ('Total de Productos', str(resumen.get('total_productos', 0))),
            ('Ventas Totales', f"${resumen.get('ventas_totales', 0):,.2f}"),
            ('Producto Más Vendido', resumen.get('producto_mas_vendido', 'N/A')),
        ]
        
        for i, (metrica, valor) in enumerate(row_data, start=1):
            cells = table_resumen.rows[i].cells
            cells[0].text = metrica
            cells[1].text = valor
            
            # Formato de celdas
            for cell in cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(10)
        
        doc.add_paragraph()  # Espacio
        
        # ============================================
        # 4. DETALLE DE PRODUCTOS
        # ============================================
        if productos:
            detalle_heading = doc.add_heading('Detalle de Productos', level=2)
            for run in detalle_heading.runs:
                run.font.color.rgb = self.colores['subtitulo']
            
            # Tabla de productos
            table_productos = doc.add_table(rows=1, cols=5)
            table_productos.style = 'Light Grid Accent 1'
            
            # Encabezados
            hdr_cells = table_productos.rows[0].cells
            headers = ['Nombre', 'Categoría', 'Precio', 'Ventas', 'Estado']
            for i, header in enumerate(headers):
                hdr_cells[i].text = header
            
            # Formato de encabezados
            for cell in hdr_cells:
                shading_elm = parse_xml(r'<w:shd {} w:fill="2980B9"/>'.format(nsdecls('w')))
                cell._tc.get_or_add_tcPr().append(shading_elm)
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True
                        run.font.color.rgb = RGBColor(255, 255, 255)
                        run.font.size = Pt(11)
            
            # Datos de productos
            for producto in productos[:50]:  # Limitar a 50 productos
                row_cells = table_productos.add_row().cells
                row_cells[0].text = producto.get('nombre', 'N/A')
                row_cells[1].text = producto.get('categoria', 'N/A')
                row_cells[2].text = f"${producto.get('precio', 0):,.2f}"
                row_cells[3].text = str(producto.get('total_ventas', 0))
                row_cells[4].text = producto.get('estado', 'N/A')
                
                # Ajustar tamaño de fuente
                for cell in row_cells:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.size = Pt(9)
            
            # Ajustar anchos de columna
            for i, width in enumerate([Inches(2.0), Inches(1.5), Inches(1.2), Inches(1.0), Inches(1.0)]):
                for row in table_productos.rows:
                    row.cells[i].width = width
        
        # ============================================
        # 5. PIE DE PÁGINA
        # ============================================
        doc.add_paragraph()
        footer = doc.add_paragraph()
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer_run = footer.add_run('_' * 80)
        footer_run.font.color.rgb = RGBColor(150, 150, 150)
        
        info_footer = doc.add_paragraph()
        info_footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        info_run = info_footer.add_run(
            f'Sistema de Reservas Turísticas | Generado el {datetime.now().strftime("%d/%m/%Y a las %H:%M")}'
        )
        info_run.font.size = Pt(8)
        info_run.font.color.rgb = RGBColor(150, 150, 150)
        info_run.italic = True
        
        # Guardar en buffer
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer
